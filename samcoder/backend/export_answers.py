#!/usr/bin/env python3
# Ekspor jawaban Prime Agent Hub — Aaron 13 Agu 2026
# Usage: python3 export_answers.py <pdf|docx|xlsx|md> <input.md> <output.ext>
import sys, re

def parse_md(md):
    lines = md.split('\n')
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            blocks.append(('h', len(m.group(1)), m.group(2)))
            i += 1; continue
        if line.startswith('```'):
            code = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code.append(lines[i]); i += 1
            i += 1
            blocks.append(('code', '\n'.join(code)))
            continue
        if line.startswith('|') and i + 1 < len(lines) and re.match(r'^\|[\s:|-]+\|?$', lines[i+1].strip()):
            rows = []
            while i < len(lines) and lines[i].startswith('|'):
                cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                rows.append(cells); i += 1
            if rows and len(rows) > 1:
                blocks.append(('table', rows))
            continue
        if re.match(r'^[-*]\s+', line) or re.match(r'^\d+\.\s+', line):
            items = []
            while i < len(lines) and (re.match(r'^[-*]\s+', lines[i].strip()) or re.match(r'^\d+\.\s+', lines[i].strip())):
                items.append(re.sub(r'^[-*]\s+|^\d+\.\s+', '', lines[i].strip()))
                i += 1
            blocks.append(('list', items))
            continue
        if line.strip() == '':
            i += 1; continue
        para = []
        while i < len(lines) and lines[i].strip() != '' and not lines[i].startswith('```') and not lines[i].startswith('|') and not re.match(r'^(#{1,6})\s+', lines[i]) and not re.match(r'^[-*]\s+', lines[i].strip()):
            para.append(lines[i].strip()); i += 1
        blocks.append(('p', ' '.join(para)))
    return blocks

def strip_md(text):
    text = re.sub(r'!\[[^\]]*\]\([^)]*\)', '', text)
    text = re.sub(r'\[([^\]]*)\]\([^)]*\)', r'\1', text)
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text

def to_pdf(md, out):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted, Table, TableStyle
    from reportlab.lib import colors
    styles = getSampleStyleSheet()
    h_styles = {1: styles['Title'], 2: styles['Heading1'], 3: styles['Heading2'], 4: styles['Heading3']}
    doc = SimpleDocTemplate(out, pagesize=A4)
    story = []
    for block in parse_md(md):
        if block[0] == 'h':
            story.append(Paragraph(strip_md(block[2]), h_styles.get(block[1], styles['Heading3'])))
        elif block[0] == 'p':
            story.append(Paragraph(strip_md(block[1]), styles['BodyText']))
        elif block[0] == 'code':
            story.append(Preformatted(block[1], styles['Code']))
        elif block[0] == 'list':
            for item in block[1]:
                story.append(Paragraph('• ' + strip_md(item), styles['BodyText']))
        elif block[0] == 'table':
            data = [[strip_md(c) for c in row] for row in block[1]]
            t = Table(data)
            t.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 0.5, colors.grey), ('BACKGROUND', (0,0), (-1,0), colors.lightgrey)]))
            story.append(t)
        story.append(Spacer(1, 6))
    doc.build(story)

def to_docx(md, out):
    from docx import Document
    doc = Document()
    for block in parse_md(md):
        if block[0] == 'h':
            doc.add_heading(strip_md(block[2]), level=min(block[1], 4))
        elif block[0] == 'p':
            doc.add_paragraph(strip_md(block[1]))
        elif block[0] == 'code':
            doc.add_paragraph(block[1], style='Intense Quote')
        elif block[0] == 'list':
            for item in block[1]:
                doc.add_paragraph(strip_md(item), style='List Bullet')
        elif block[0] == 'table':
            data = [[strip_md(c) for c in row] for row in block[1]]
            if data:
                table = doc.add_table(rows=len(data), cols=len(data[0]))
                for ri, row in enumerate(data):
                    for ci, val in enumerate(row):
                        table.cell(ri, ci).text = val
    doc.save(out)

def to_xlsx(md, out):
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = 'Jawaban'
    tables = [b[1] for b in parse_md(md) if b[0] == 'table']
    if tables:
        for row in tables[0]:
            ws.append([strip_md(c) for c in row])
    ws2 = wb.create_sheet('Ringkasan')
    for b in parse_md(md):
        if b[0] == 'h':
            ws2.append([strip_md(b[2])])
        elif b[0] == 'p':
            ws2.append([strip_md(b[1])])
    wb.save(out)

def to_md(md, out):
    with open(out, 'w', encoding='utf-8') as f:
        f.write(md)

if __name__ == '__main__':
    fmt = sys.argv[1]
    inp = sys.argv[2]
    outp = sys.argv[3]
    with open(inp, 'r', encoding='utf-8') as f:
        md = f.read()
    {'pdf': to_pdf, 'docx': to_docx, 'xlsx': to_xlsx, 'md': to_md}[fmt](md, outp)
